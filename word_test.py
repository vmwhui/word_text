import docx

#导入模板文档
document = docx.Document("template.docx")

#读取文档函数
def read_text():
    #读取模板文档所有段落
    all_paragraphs = document.paragraphs
    #开始循环段落
    i = 1
    for paragraphs in all_paragraphs:
        print("第" + str(i) + str(paragraphs.text))
        i = i + 1
    #读取模板文档所有表格
    cell_vl = []
    which_tab = 1
    all_tables = document.tables
    all_tables[3].add_row()
    for table in all_tables:
        print("第%s张表" % which_tab)
        for row in table.rows:
            cell_vl.clear()
            for cell in row.cells:
                cell_vl.append(cell.text)
            print(cell_vl)
        which_tab = which_tab + 1
read_text()

document.save("A1.docx")